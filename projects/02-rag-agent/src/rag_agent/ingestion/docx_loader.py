"""Word (DOCX) document loader."""

from __future__ import annotations

from pathlib import Path

from rag_agent.ingestion.base import BaseLoader, Document
from rag_agent.utils.logger import get_logger

logger = get_logger()


class DocxLoader(BaseLoader):
    """Loads Microsoft Word (.docx) files."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def load(self, source: str) -> list[Document]:
        path = Path(source).resolve()

        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))

            # Extract text from paragraphs
            paragraphs: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            content = "\n\n".join(paragraphs)

            if not content.strip():
                logger.warning(f"No text extracted from {source}")
                content = f"[No extractable text from {path.name}]"

            logger.debug(f"Loaded {len(paragraphs)} paragraphs from {source}")

            return [
                Document(
                    content=content,
                    metadata={
                        "source_path": str(path),
                        "file_name": path.name,
                        "file_type": "docx",
                    },
                )
            ]

        except Exception as e:
            logger.error(f"Failed to load DOCX {source}: {e}")
            raise
